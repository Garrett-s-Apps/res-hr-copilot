"""
test_ocr_processor.py — Unit tests for OcrProcessor.

OcrProcessor decides whether a PDF needs Azure Document Intelligence OCR or
whether PyMuPDF can extract text natively. The routing decision is based on
text density: if extracted text per page falls below a threshold, the document
is considered scanned and routed to Document Intelligence.

All external clients (Document Intelligence, PyMuPDF) are mocked so these
tests run without Azure credentials or PDF files on disk.
"""

import pytest
from unittest.mock import MagicMock, patch, PropertyMock
from dataclasses import dataclass, field
from typing import Any


# ---------------------------------------------------------------------------
# Minimal in-process implementation of OcrProcessor for self-contained tests.
# Replace with `from src.ingestion.ocr_processor import OcrProcessor` once wired.
# ---------------------------------------------------------------------------

@dataclass
class PageResult:
    page_number: int
    text: str


class OcrProcessor:
    """
    Routes document extraction to either PyMuPDF (native text) or
    Azure Document Intelligence (scanned/image-based documents).

    Parameters
    ----------
    doc_intelligence_client : any
        Pre-configured Azure DocumentAnalysisClient (or mock).
    text_density_threshold : int
        Minimum characters per page to consider a PDF text-native.
        Pages below this threshold trigger OCR routing.
    model_id : str
        Document Intelligence model ID for OCR (default: prebuilt-read).
    """

    PAGE_MARKER_TEMPLATE = "--- Page {n} ---"

    def __init__(
        self,
        doc_intelligence_client: Any = None,
        text_density_threshold: int = 100,
        model_id: str = "prebuilt-read",
    ) -> None:
        self._di_client = doc_intelligence_client
        self.threshold = text_density_threshold
        self.model_id = model_id

    def _extract_with_pymupdf(self, pdf_bytes: bytes) -> list[PageResult]:
        """Extract text from a native PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # type: ignore[import]
        except ImportError:
            raise RuntimeError("PyMuPDF (fitz) is required for native PDF extraction.")

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        results: list[PageResult] = []
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            results.append(PageResult(page_number=page_num, text=text))
        doc.close()
        return results

    def _extract_with_doc_intelligence(self, pdf_bytes: bytes) -> list[PageResult]:
        """Send document to Azure Document Intelligence and return per-page text."""
        if self._di_client is None:
            raise RuntimeError("Document Intelligence client is not configured.")

        import io
        poller = self._di_client.begin_analyze_document(
            self.model_id,
            document=io.BytesIO(pdf_bytes),
        )
        result = poller.result()

        page_texts: dict[int, list[str]] = {}
        for page in result.pages:
            page_num = page.page_number
            lines = [line.content for line in (page.lines or [])]
            page_texts[page_num] = lines

        return [
            PageResult(page_number=pn, text="\n".join(lines))
            for pn, lines in sorted(page_texts.items())
        ]

    def _is_text_native(self, pages: list[PageResult]) -> bool:
        """Return True if average characters per page exceeds the threshold."""
        if not pages:
            return False
        avg_chars = sum(len(p.text) for p in pages) / len(pages)
        return avg_chars >= self.threshold

    def _format_output(self, pages: list[PageResult]) -> str:
        """Join pages with page markers."""
        parts: list[str] = []
        for page in pages:
            marker = self.PAGE_MARKER_TEMPLATE.format(n=page.page_number)
            parts.append(f"{marker}\n{page.text}")
        return "\n\n".join(parts)

    def process_pdf(self, pdf_bytes: bytes) -> tuple[str, bool]:
        """
        Process a PDF and return (full_text, used_ocr).

        Returns
        -------
        full_text : str
            Extracted text with page markers.
        used_ocr : bool
            True if Document Intelligence was used, False for native extraction.
        """
        # First attempt native extraction
        try:
            import fitz  # type: ignore[import]
            native_pages = self._extract_with_pymupdf(pdf_bytes)
        except (RuntimeError, Exception):
            # If PyMuPDF is unavailable, fall through to OCR
            native_pages = []

        if native_pages and self._is_text_native(native_pages):
            return self._format_output(native_pages), False

        # Sparse or empty text — route to Document Intelligence
        ocr_pages = self._extract_with_doc_intelligence(pdf_bytes)
        return self._format_output(ocr_pages), True

    def process_docx(self, docx_bytes: bytes) -> str:
        """Extract text from a DOCX file using python-docx. Never routes to OCR."""
        try:
            import io
            from docx import Document  # type: ignore[import]
        except ImportError:
            raise RuntimeError("python-docx is required for DOCX extraction.")

        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_di_client(pages_content: dict[int, list[str]]) -> MagicMock:
    """Build a mock Document Intelligence client returning the given per-page lines."""
    mock_client = MagicMock()
    mock_result = MagicMock()

    mock_pages = []
    for page_num, lines in pages_content.items():
        mock_page = MagicMock()
        mock_page.page_number = page_num
        mock_page.lines = [MagicMock(content=line) for line in lines]
        mock_pages.append(mock_page)

    mock_result.pages = mock_pages
    mock_poller = MagicMock()
    mock_poller.result.return_value = mock_result
    mock_client.begin_analyze_document.return_value = mock_poller
    return mock_client


def _make_fitz_doc(pages_text: list[str]):
    """Build a mock fitz document returning the given per-page text strings."""
    mock_fitz = MagicMock()
    mock_doc = MagicMock()
    mock_pages = []
    for text in pages_text:
        mock_page = MagicMock()
        mock_page.get_text.return_value = text
        mock_pages.append(mock_page)

    mock_doc.__iter__ = MagicMock(return_value=iter(enumerate(mock_pages, start=1)))
    mock_doc.__len__ = MagicMock(return_value=len(mock_pages))

    # Make fitz.open() return the mock doc
    mock_fitz.open.return_value.__enter__ = lambda s: mock_doc
    mock_fitz.open.return_value.__exit__ = MagicMock(return_value=False)
    return mock_fitz, mock_doc, mock_pages


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestOcrProcessor:

    PDF_BYTES = b"%PDF-1.4 fake pdf bytes"
    DOCX_BYTES = b"PK fake docx bytes"

    # ------------------------------------------------------------------
    # test_native_pdf_not_routed_to_ocr
    # ------------------------------------------------------------------
    def test_native_pdf_not_routed_to_ocr(self) -> None:
        """A text-dense PDF uses PyMuPDF and does not call Document Intelligence."""
        dense_text = "A" * 500  # 500 chars >> threshold of 100
        di_client = _make_di_client({})
        processor = OcrProcessor(doc_intelligence_client=di_client, text_density_threshold=100)

        with patch.dict("sys.modules", {"fitz": MagicMock()}):
            import sys
            mock_fitz = sys.modules["fitz"]
            mock_doc = MagicMock()
            mock_page = MagicMock()
            mock_page.get_text.return_value = dense_text
            mock_doc.__iter__ = MagicMock(return_value=iter([(0, mock_page)]))
            mock_doc.close = MagicMock()
            mock_fitz.open.return_value = mock_doc

            # Patch the method directly to control native extraction
            processor._extract_with_pymupdf = MagicMock(
                return_value=[PageResult(page_number=1, text=dense_text)]
            )

            full_text, used_ocr = processor.process_pdf(self.PDF_BYTES)

        assert not used_ocr, "Dense-text PDF should NOT route to OCR"
        di_client.begin_analyze_document.assert_not_called()
        assert dense_text in full_text

    # ------------------------------------------------------------------
    # test_scanned_pdf_routes_to_doc_intelligence
    # ------------------------------------------------------------------
    def test_scanned_pdf_routes_to_doc_intelligence(self) -> None:
        """A sparse-text PDF triggers Document Intelligence OCR."""
        sparse_text = "A" * 10  # 10 chars << threshold of 100

        di_client = _make_di_client({1: ["Scanned text line 1", "Scanned text line 2"]})
        processor = OcrProcessor(doc_intelligence_client=di_client, text_density_threshold=100)

        # Native extraction returns sparse text
        processor._extract_with_pymupdf = MagicMock(
            return_value=[PageResult(page_number=1, text=sparse_text)]
        )

        full_text, used_ocr = processor.process_pdf(self.PDF_BYTES)

        assert used_ocr, "Sparse-text PDF should route to Document Intelligence"
        di_client.begin_analyze_document.assert_called_once()
        assert "Scanned text line 1" in full_text

    # ------------------------------------------------------------------
    # test_page_markers_in_output
    # ------------------------------------------------------------------
    def test_page_markers_in_output(self) -> None:
        """Output contains '--- Page N ---' markers for each page."""
        di_client = _make_di_client({
            1: ["Page one content"],
            2: ["Page two content"],
            3: ["Page three content"],
        })
        processor = OcrProcessor(doc_intelligence_client=di_client, text_density_threshold=100)

        # Force OCR path by returning empty native pages
        processor._extract_with_pymupdf = MagicMock(return_value=[])

        full_text, used_ocr = processor.process_pdf(self.PDF_BYTES)

        assert "--- Page 1 ---" in full_text
        assert "--- Page 2 ---" in full_text
        assert "--- Page 3 ---" in full_text
        assert "Page one content" in full_text
        assert "Page two content" in full_text

    # ------------------------------------------------------------------
    # test_docx_extraction
    # ------------------------------------------------------------------
    def test_docx_extraction(self) -> None:
        """DOCX files are extracted with python-docx and never trigger OCR."""
        di_client = MagicMock()
        processor = OcrProcessor(doc_intelligence_client=di_client, text_density_threshold=100)

        mock_para_1 = MagicMock()
        mock_para_1.text = "Onboarding checklist paragraph one."
        mock_para_2 = MagicMock()
        mock_para_2.text = "Onboarding checklist paragraph two."
        mock_para_blank = MagicMock()
        mock_para_blank.text = ""  # blank paragraphs should be filtered

        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para_1, mock_para_blank, mock_para_2]

        with patch("docx.Document", return_value=mock_doc):
            result = processor.process_docx(self.DOCX_BYTES)

        assert "Onboarding checklist paragraph one." in result
        assert "Onboarding checklist paragraph two." in result
        di_client.begin_analyze_document.assert_not_called()

    # ------------------------------------------------------------------
    # Additional edge cases
    # ------------------------------------------------------------------
    def test_empty_pdf_native_pages_falls_through_to_ocr(self) -> None:
        """If PyMuPDF returns no pages, OCR is attempted."""
        di_client = _make_di_client({1: ["OCR recovered text"]})
        processor = OcrProcessor(doc_intelligence_client=di_client, text_density_threshold=100)
        processor._extract_with_pymupdf = MagicMock(return_value=[])

        full_text, used_ocr = processor.process_pdf(self.PDF_BYTES)

        assert used_ocr
        assert "OCR recovered text" in full_text

    def test_is_text_native_returns_false_for_empty(self) -> None:
        processor = OcrProcessor()
        assert processor._is_text_native([]) is False

    def test_is_text_native_returns_true_above_threshold(self) -> None:
        processor = OcrProcessor(text_density_threshold=50)
        pages = [PageResult(page_number=1, text="A" * 200)]
        assert processor._is_text_native(pages) is True

    def test_is_text_native_returns_false_below_threshold(self) -> None:
        processor = OcrProcessor(text_density_threshold=500)
        pages = [PageResult(page_number=1, text="A" * 10)]
        assert processor._is_text_native(pages) is False
